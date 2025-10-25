"""
Document processing and chunking functionality for RAG system
"""
import os
import logging
from typing import List, Dict, Any
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import markdown

# LangChain components
try:
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_core.documents import Document as LangChainDocument
except ImportError:
    # Fallback for older versions
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangChainDocument

from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document loading, processing, and chunking"""
    
    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n\n", "\n\n", "\n", ". ", " ", ""]  # Better separators for markdown
        )
    
    def load_document(self, file_path: str) -> str:
        """Load document content based on file extension"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._load_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return self._load_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._load_txt(file_path)
            elif file_path.suffix.lower() == '.md':
                return self._load_markdown(file_path)
            elif file_path.suffix.lower() == '.html':
                return self._load_html(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: Path) -> str:
        """Load PDF document"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _load_docx(self, file_path: Path) -> str:
        """Load DOCX document"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _load_txt(self, file_path: Path) -> str:
        """Load text document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _load_markdown(self, file_path: Path) -> str:
        """Load markdown document with better formatting"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            
        # Process markdown to keep headers with content
        lines = md_content.split('\n')
        processed_lines = []
        
        for i, line in enumerate(lines):
            # If this is a header (starts with #), keep it with the next few lines
            if line.strip().startswith('#'):
                processed_lines.append(line)
                # Add the next few lines that aren't headers
                j = i + 1
                while j < len(lines) and j < i + 10:  # Look ahead up to 10 lines
                    next_line = lines[j].strip()
                    if next_line and not next_line.startswith('#'):
                        processed_lines.append(lines[j])
                    elif next_line.startswith('#'):
                        break
                    j += 1
            else:
                processed_lines.append(line)
        
        return '\n'.join(processed_lines)
    
    def _load_html(self, file_path: Path) -> str:
        """Load HTML document"""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
            soup = BeautifulSoup(html_content, 'html.parser')
            return soup.get_text()
    
    def chunk_document(self, text: str, metadata: Dict[str, Any] = None) -> List[LangChainDocument]:
        """Split document into chunks"""
        if metadata is None:
            metadata = {}
        
        # Create LangChain document
        doc = LangChainDocument(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        logger.info(f"Document split into {len(chunks)} chunks")
        return chunks
    
    def process_file(self, file_path: str) -> List[LangChainDocument]:
        """Complete processing pipeline for a single file"""
        logger.info(f"Processing file: {file_path}")
        
        # Load document
        text = self.load_document(file_path)
        
        # Create metadata
        file_path_obj = Path(file_path)
        metadata = {
            "source": str(file_path),
            "filename": file_path_obj.name,
            "file_type": file_path_obj.suffix,
            "file_size": file_path_obj.stat().st_size
        }
        
        # Chunk document
        chunks = self.chunk_document(text, metadata)
        
        return chunks
    
    def process_directory(self, directory_path: str) -> List[LangChainDocument]:
        """Process all supported files in a directory"""
        directory = Path(directory_path)
        all_chunks = []
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    chunks = self.process_file(str(file_path))
                    all_chunks.extend(chunks)
                except Exception as e:
                    logger.error(f"Error processing {file_path}: {str(e)}")
                    continue
        
        logger.info(f"Processed {len(all_chunks)} total chunks from directory")
        return all_chunks


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    # Process a single file
    # chunks = processor.process_file("sample_document.pdf")
    
    # Process a directory
    # chunks = processor.process_directory("./documents")
    
    print("Document processor ready!")
